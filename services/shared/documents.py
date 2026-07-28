import csv
import hashlib
import io
import json
import re
import unicodedata
import zipfile
from pathlib import PurePath

from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_DOCUMENT_BYTES = 10_485_760
MAX_ARCHIVE_MEMBERS = 2_000
MAX_UNCOMPRESSED_BYTES = 50_000_000
SECRET_PATTERN = re.compile(
    r"(?i)(api[_-]?key|secret|password|token)\s*[:=]\s*[\"']?[A-Za-z0-9_\-]{8,}"
)
FORMULA_PREFIXES = ("=", "+", "-", "@")
ALLOWED = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".csv": "text/csv", ".html": "text/html", ".htm": "text/html", ".txt": "text/plain",
    ".md": "text/markdown", ".json": "application/json",
}


class UnsafeDocument(ValueError):
    pass


def safe_filename(name):
    base = PurePath(name.replace("\\", "/")).name
    normalized = unicodedata.normalize("NFKC", base)
    safe = re.sub(r"[^A-Za-z0-9._-]+", "_", normalized).strip("._")
    if not safe or safe != normalized or "/" in name or "\\" in name or ".." in name:
        raise UnsafeDocument("nombre de archivo hostil o traversal")
    return safe


def _check_archive(data):
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        infos = archive.infolist()
        total = sum(item.file_size for item in infos)
        if len(infos) > MAX_ARCHIVE_MEMBERS or total > MAX_UNCOMPRESSED_BYTES:
            raise UnsafeDocument("archivo ZIP sobredimensionado")
        if any("vbaProject.bin" in item.filename or item.filename.startswith(("/", "\\")) or ".." in PurePath(item.filename).parts for item in infos):
            raise UnsafeDocument("macro o ruta insegura dentro del archivo")


def _extract_pdf(data):
    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        raise UnsafeDocument("PDF cifrado no admitido")
    return [(index + 1, page.extract_text() or "") for index, page in enumerate(reader.pages)]


def _extract_docx(data):
    document = Document(io.BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return [(None, text)]


def _extract_xlsx(data):
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=False, keep_links=False)
    chunks = []
    for sheet in workbook.worksheets:
        lines = []
        for row in sheet.iter_rows(values_only=True):
            values = []
            for value in row:
                rendered = "" if value is None else str(value)
                if rendered.startswith(FORMULA_PREFIXES):
                    rendered = "'" + rendered
                values.append(rendered)
            lines.append("\t".join(values))
        chunks.append((None, f"[Hoja: {sheet.title}]\n" + "\n".join(lines)))
    return chunks


def _extract_text(data, extension):
    text = data.decode("utf-8-sig")
    if extension in {".html", ".htm"}:
        soup = BeautifulSoup(text, "lxml")
        for element in soup(["script", "style", "iframe", "object", "embed", "form"]):
            element.decompose()
        for tag in soup.find_all(True):
            for attribute in list(tag.attrs):
                if attribute.lower().startswith("on") or attribute.lower() in {"src", "href"}:
                    del tag.attrs[attribute]
        text = soup.get_text("\n")
    elif extension == ".json":
        text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
    elif extension == ".csv":
        output = io.StringIO()
        writer = csv.writer(output)
        for row in csv.reader(io.StringIO(text)):
            writer.writerow([("'" + cell) if cell.startswith(FORMULA_PREFIXES) else cell for cell in row])
        text = output.getvalue()
    return [(None, text)]


def inspect_and_extract(name, media_type, data):
    if len(data) > MAX_DOCUMENT_BYTES:
        raise UnsafeDocument("archivo sobredimensionado")
    safe = safe_filename(name)
    extension = PurePath(safe).suffix.lower()
    expected = ALLOWED.get(extension)
    if not expected:
        raise UnsafeDocument("formato no permitido")
    if media_type.split(";")[0].strip().lower() != expected:
        raise UnsafeDocument("MIME no coincide con la extensión")
    if extension == ".pdf" and not data.startswith(b"%PDF"):
        raise UnsafeDocument("MIME PDF falso")
    if extension in {".docx", ".xlsx"}:
        if not data.startswith(b"PK"):
            raise UnsafeDocument("MIME Office falso")
        _check_archive(data)
    extractor = _extract_pdf if extension == ".pdf" else _extract_docx if extension == ".docx" else _extract_xlsx if extension == ".xlsx" else None
    chunks = extractor(data) if extractor else _extract_text(data, extension)
    findings = []
    fragments = []
    ordinal = 0
    for page, text_value in chunks:
        clean = "\n".join(line.strip() for line in text_value.splitlines() if line.strip())
        if not clean:
            continue
        if SECRET_PATTERN.search(clean):
            findings.append("posible secreto redactado")
            clean = SECRET_PATTERN.sub("[SECRETO REDACTADO]", clean)
        if re.search(r"(?i)(ignore|ignora).{0,30}(instruction|instrucci)", clean):
            findings.append("instrucción documental tratada como dato")
        if re.search(r"https?://", clean):
            findings.append("enlace externo conservado como texto no ejecutable")
        for start in range(0, len(clean), 2_000):
            fragment = clean[start:start + 2_000]
            fragments.append({"ordinal": ordinal, "page": page, "line_start": clean[:start].count("\n") + 1,
                              "line_end": clean[:start + len(fragment)].count("\n") + 1,
                              "text": fragment, "sha256": hashlib.sha256(fragment.encode("utf-8")).hexdigest()})
            ordinal += 1
    return {"safe_name": safe, "media_type": expected, "size_bytes": len(data),
            "sha256": hashlib.sha256(data).hexdigest(), "security_findings": sorted(set(findings)),
            "fragments": fragments, "result_kind": "observed"}
